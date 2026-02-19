"""Deterministic checks — whitespace, enumerations, references."""

from __future__ import annotations

import os
import re
from typing import Any

from docx import Document

from mcp_server.docx_parser import _detect_heading_level


# ── whitespace ─────────────────────────────────────────────────────

_RE_DOUBLE_SPACE = re.compile(r"  +")


def _find_parent_heading(paragraphs, idx: int) -> str | None:
    """Walk backwards from idx to find the nearest heading."""
    for i in range(idx - 1, -1, -1):
        level = _detect_heading_level(paragraphs[i])
        if level is not None:
            return paragraphs[i].text.strip()
    return None


def check_whitespace(filepath: str) -> dict[str, Any]:
    """Find whitespace issues in a .docx document.

    Checks for:
      - Double (or more) consecutive spaces within text
      - Trailing whitespace at end of paragraph text
      - Leading whitespace at start of paragraph text
      - Consecutive blank paragraphs

    Returns:
      {filepath, issue_count, issues: [{type, paragraph_index,
       section, text, detail}, ...]}
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    paragraphs = doc.paragraphs
    issues: list[dict[str, Any]] = []

    prev_blank = False

    for idx, p in enumerate(paragraphs):
        text = p.text
        level = _detect_heading_level(p)

        # Skip headings themselves
        if level is not None:
            prev_blank = False
            continue

        section = _find_parent_heading(paragraphs, idx)
        is_blank = text.strip() == ""

        # Consecutive blank paragraphs
        if is_blank and prev_blank:
            issues.append({
                "type": "consecutive_blank_paragraphs",
                "paragraph_index": idx,
                "section": section,
                "text": "",
                "detail": "Multiple consecutive blank paragraphs",
            })
        prev_blank = is_blank

        if is_blank:
            continue

        # Double spaces
        for m in _RE_DOUBLE_SPACE.finditer(text):
            start = max(0, m.start() - 20)
            end = min(len(text), m.end() + 20)
            context = text[start:end]
            issues.append({
                "type": "double_space",
                "paragraph_index": idx,
                "section": section,
                "text": text[:80] + ("…" if len(text) > 80 else ""),
                "detail": f"Multiple spaces at position {m.start()}: \"…{context}…\"",
            })

        # Trailing whitespace
        if text != text.rstrip():
            trailing = text[len(text.rstrip()):]
            issues.append({
                "type": "trailing_whitespace",
                "paragraph_index": idx,
                "section": section,
                "text": text.rstrip()[:80] + ("…" if len(text.rstrip()) > 80 else ""),
                "detail": f"{len(trailing)} trailing whitespace character(s)",
            })

        # Leading whitespace (skip list items — they may be indented)
        if text != text.lstrip() and not _is_list_item(p):
            leading = text[:len(text) - len(text.lstrip())]
            issues.append({
                "type": "leading_whitespace",
                "paragraph_index": idx,
                "section": section,
                "text": text.strip()[:80] + ("…" if len(text.strip()) > 80 else ""),
                "detail": f"{len(leading)} leading whitespace character(s)",
            })

    return {
        "filepath": filepath,
        "issue_count": len(issues),
        "issues": issues,
    }


def _is_list_item(paragraph) -> bool:
    """Check if a paragraph is a Word list item (has numPr in XML)."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pPr = paragraph._element.find(f"{{{W}}}pPr")
    if pPr is not None:
        numPr = pPr.find(f"{{{W}}}numPr")
        if numPr is not None:
            return True
    return False


# ── enumeration check ──────────────────────────────────────────────

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_RE_ENUM_BOTH_PARENS = re.compile(r'^\(([a-z]{1,4})\)\s', re.IGNORECASE)
_RE_ENUM_RIGHT_PAREN = re.compile(r'^([a-z]{1,4})\)\s', re.IGNORECASE)


def _get_numPr(paragraph) -> tuple[int, int] | None:
    """Extract (numId, ilvl) from paragraph XML, or None if not a Word list item."""
    pPr = paragraph._element.find(f"{{{_W}}}pPr")
    if pPr is None:
        return None
    numPr = pPr.find(f"{{{_W}}}numPr")
    if numPr is None:
        return None
    numId_el = numPr.find(f"{{{_W}}}numId")
    ilvl_el = numPr.find(f"{{{_W}}}ilvl")
    if numId_el is None or ilvl_el is None:
        return None
    numId = numId_el.get(f"{{{_W}}}val")
    ilvl = ilvl_el.get(f"{{{_W}}}val")
    if numId is None or ilvl is None:
        return None
    return (int(numId), int(ilvl))


def _get_num_format(doc, numId: int, ilvl: int) -> str | None:
    """Get the numbering format string (e.g. 'lowerLetter') for a numId/ilvl pair."""
    try:
        numbering_part = doc.part.numbering_part
    except AttributeError:
        return None
    if numbering_part is None:
        return None
    root = numbering_part._element
    abs_num_id = None
    for num_el in root.findall(f"{{{_W}}}num"):
        if num_el.get(f"{{{_W}}}numId") == str(numId):
            abs_ref = num_el.find(f"{{{_W}}}abstractNumId")
            if abs_ref is not None:
                abs_num_id = abs_ref.get(f"{{{_W}}}val")
            break
    if abs_num_id is None:
        return None
    for abs_num in root.findall(f"{{{_W}}}abstractNum"):
        if abs_num.get(f"{{{_W}}}abstractNumId") == abs_num_id:
            for lvl in abs_num.findall(f"{{{_W}}}lvl"):
                if lvl.get(f"{{{_W}}}ilvl") == str(ilvl):
                    num_fmt = lvl.find(f"{{{_W}}}numFmt")
                    if num_fmt is not None:
                        return num_fmt.get(f"{{{_W}}}val")
    return None


def _detect_text_list_pattern(text: str) -> dict[str, str] | None:
    """Detect enumeration marker at the start of text.

    Matches patterns like: (a), (ii), a), iii)

    Returns dict with 'style' ('(x)' or 'x)') and 'marker' (str),
    or None if no pattern matched.
    """
    m = _RE_ENUM_BOTH_PARENS.match(text)
    if m:
        return {"style": "(x)", "marker": m.group(1).lower()}
    m = _RE_ENUM_RIGHT_PAREN.match(text)
    if m:
        return {"style": "x)", "marker": m.group(1).lower()}
    return None


def _get_terminator(text: str) -> str:
    """Return the last non-space character of text, or '' for blank/empty text."""
    stripped = text.rstrip()
    return stripped[-1] if stripped else ""


def _check_list_delimiters(items: list[dict]) -> list[dict[str, Any]]:
    """Check delimiter consistency within an enumeration run.

    Flags when non-last items use mixed terminators (e.g. mixing ',' and ';').

    Returns a (possibly empty) list of issue description dicts.
    """
    if len(items) < 2:
        return []

    terminators = [_get_terminator(it["text"]) for it in items]
    non_last_terms = terminators[:-1]
    relevant = {t for t in non_last_terms if t in {",", ";", "."}}

    if len(relevant) > 1:
        term_str = "/".join(terminators)
        return [{
            "check": "terminator_inconsistency",
            "detail": f"Inconsistent terminators among enumeration items: {term_str}",
            "terminators": terminators,
        }]
    return []


def check_enumerations(filepath: str) -> dict[str, Any]:
    """Check enumeration delimiter consistency in a .docx document.

    Detects text-pattern list items: (a)/(b), a)/b), (i)/(ii), etc.
    Reports runs where non-last items use mixed terminators (e.g. ',' and ';').

    Returns:
      {filepath, issue_count, issues: [{type, paragraph_index, section,
       text, detail, terminators}, ...]}
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    paragraphs = doc.paragraphs
    result_issues: list[dict[str, Any]] = []

    current_run: list[dict[str, Any]] = []
    current_section: str | None = None

    def flush() -> None:
        for prob in _check_list_delimiters(current_run):
            first = current_run[0]
            result_issues.append({
                "type": prob["check"],
                "paragraph_index": first["paragraph_index"],
                "section": first["section"],
                "text": first["text"][:80] + ("…" if len(first["text"]) > 80 else ""),
                "detail": prob["detail"],
                "terminators": prob["terminators"],
            })

    for idx, p in enumerate(paragraphs):
        level = _detect_heading_level(p)
        if level is not None:
            flush()
            current_run = []
            current_section = p.text.strip()
            continue

        text = p.text
        if not text.strip():
            flush()
            current_run = []
            continue

        item_info = _detect_text_list_pattern(text)
        if item_info is None:
            flush()
            current_run = []
            continue

        # Flush and restart if delimiter style changes mid-run
        if current_run and current_run[-1]["style"] != item_info["style"]:
            flush()
            current_run = []

        current_run.append({
            "paragraph_index": idx,
            "section": current_section,
            "text": text,
            "style": item_info["style"],
            "marker": item_info["marker"],
        })

    flush()  # handle any trailing run at end of document

    return {
        "filepath": filepath,
        "issue_count": len(result_issues),
        "issues": result_issues,
    }


# ── reference extraction & validation ──────────────────────────────

# Text extraction patterns — various Czech grammatical forms
_RE_CL_ABBR  = re.compile(r'\bčl\.\s*(\d+(?:\.\d+)*)\b')
_RE_CLANEK   = re.compile(r'\bčlánk\w{0,3}\s+(\d+(?:\.\d+)*)\b', re.IGNORECASE)
_RE_PRILOHA  = re.compile(r'\bpřílo\w{1,4}\s+č\.\s*(\d+)\b', re.IGNORECASE)
_RE_PARA_LAW = re.compile(r'§\s*(\d+\w*)\b')

# Heading match patterns for building the valid-target sets
_RE_H_CLANEK  = re.compile(r'^[Čč]lánek\s+(\d+)\b')
_RE_H_PRILOHA = re.compile(r'^[Pp]říloha\s+č\.\s*(\d+)\b')


def _extract_field_codes(doc) -> list[dict[str, Any]]:
    """Walk paragraph XML to extract REF/PAGEREF Word field codes.

    Returns list of {instr, display_text, paragraph_index}.
    """
    result: list[dict[str, Any]] = []

    for p_idx, para in enumerate(doc.paragraphs):
        state: str | None = None
        instr_buf: list[str] = []
        display_buf: list[str] = []

        for run_el in para._element:
            if run_el.tag != f"{{{_W}}}r":
                continue
            for child in run_el:
                local = child.tag.split("}")[1] if "}" in child.tag else child.tag

                if local == "fldChar":
                    fc_type = child.get(f"{{{_W}}}fldCharType")
                    if fc_type == "begin":
                        state, instr_buf, display_buf = "instr", [], []
                    elif fc_type == "separate":
                        state = "display"
                    elif fc_type == "end":
                        instr = "".join(instr_buf).strip()
                        if instr.upper().lstrip().startswith(("REF ", "PAGEREF ")):
                            result.append({
                                "instr": instr,
                                "display_text": "".join(display_buf).strip(),
                                "paragraph_index": p_idx,
                            })
                        state = None
                elif local == "instrText" and state == "instr":
                    instr_buf.append(child.text or "")
                elif local == "t" and state == "display":
                    display_buf.append(child.text or "")

    return result


def _extract_text_references(paragraphs) -> list[dict[str, Any]]:
    """Extract Czech legal text references from paragraphs.

    Returns list of {text, raw, type, target, section, paragraph_index}.
    """
    result: list[dict[str, Any]] = []
    current_section: str | None = None

    for idx, para in enumerate(paragraphs):
        level = _detect_heading_level(para)
        if level is not None:
            current_section = para.text.strip()
            continue

        text = para.text
        if not text.strip():
            continue

        def _add(m, ref_type: str, normalized: str, target: str) -> None:
            result.append({
                "text": normalized,
                "raw": m.group(0).strip(),
                "type": ref_type,
                "target": target,
                "section": current_section,
                "paragraph_index": idx,
            })

        for m in _RE_CL_ABBR.finditer(text):
            _add(m, "článek", f"článek {m.group(1)}", m.group(1))

        for m in _RE_CLANEK.finditer(text):
            _add(m, "článek", f"článek {m.group(1)}", m.group(1))

        for m in _RE_PRILOHA.finditer(text):
            _add(m, "příloha", f"příloha č. {m.group(1)}", m.group(1))

        for m in _RE_PARA_LAW.finditer(text):
            _add(m, "§", f"§ {m.group(1)}", m.group(1))

    return result


def _get_bookmarks(doc) -> list[str]:
    """Extract all bookmark names from the document."""
    bookmarks: list[str] = []
    for para in doc.paragraphs:
        for el in para._element.iter(f"{{{_W}}}bookmarkStart"):
            name = el.get(f"{{{_W}}}name")
            if name:
                bookmarks.append(name)
    return bookmarks


def _validate_references(
    refs: list[dict[str, Any]],
    headings: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Split refs into (valid, invalid) using headings as the target set.

    Only validates 'článek' and 'příloha' refs. § and others pass as valid.
    """
    article_nums: set[str] = set()
    annex_nums: set[str] = set()
    for h in headings:
        m = _RE_H_CLANEK.match(h["text"])
        if m:
            article_nums.add(m.group(1))
            continue
        m = _RE_H_PRILOHA.match(h["text"])
        if m:
            annex_nums.add(m.group(1))

    valid: list[dict[str, Any]] = []
    invalid: list[dict[str, Any]] = []

    for ref in refs:
        if ref["type"] == "článek":
            base = ref["target"].split(".")[0]
            (valid if base in article_nums else invalid).append(ref)
        elif ref["type"] == "příloha":
            (valid if ref["target"] in annex_nums else invalid).append(ref)
        else:
            valid.append(ref)  # § — external law citations, no validation

    return valid, invalid


def extract_and_validate_references(filepath: str) -> dict[str, Any]:
    """Extract and validate all cross-references in a .docx document.

    Detects Word field codes (REF/PAGEREF) and plain text Czech legal
    references (čl., článek, příloha č., §).

    Validates article and annex text refs against document headings.
    All internal text refs (article + annex) are reported as
    field_code_violations — they should use Word REF fields in a properly
    formatted legal document.

    Returns:
      {filepath, all_refs, valid, invalid, field_code_refs,
       field_code_violations, bookmarks}
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    paragraphs = doc.paragraphs

    headings = [
        {"level": _detect_heading_level(p), "text": p.text.strip()}
        for p in paragraphs
        if _detect_heading_level(p) is not None
    ]

    field_code_refs = _extract_field_codes(doc)
    text_refs = _extract_text_references(paragraphs)
    bookmarks = _get_bookmarks(doc)
    valid_refs, invalid_refs = _validate_references(text_refs, headings)

    field_code_violations = [
        r for r in text_refs if r["type"] in ("článek", "příloha")
    ]

    return {
        "filepath": filepath,
        "all_refs": text_refs + [{**r, "type": "field_code"} for r in field_code_refs],
        "valid": valid_refs,
        "invalid": invalid_refs,
        "field_code_refs": field_code_refs,
        "field_code_violations": field_code_violations,
        "bookmarks": bookmarks,
    }
