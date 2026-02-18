"""Document parsing — heading tree, section content, summaries.

Provides three public functions (exposed as MCP tools in server.py):
  load_document_structure  — heading tree + metadata
  get_section_content      — full text under a heading
  get_all_sections_summary — compact preview + MD5 hash per section
"""

from __future__ import annotations

import hashlib
import os
from typing import Any

from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── helpers ────────────────────────────────────────────────────────


def _detect_heading_level(paragraph) -> int | None:
    """Return heading level (0-9) or None if paragraph is not a heading.

    Strategy:
      1. Style name like "Heading 1" → level 1
      2. XML fallback: <w:pPr><w:outlineLvl w:val="0"/> → level 1
    """
    style_name = paragraph.style.name or ""

    # "Heading 1", "Heading 2", ... or "Title" (level 0)
    if style_name == "Title":
        return 0
    if style_name.startswith("Heading"):
        parts = style_name.split()
        if len(parts) == 2 and parts[1].isdigit():
            return int(parts[1])

    # XML fallback — custom styles with outline level set
    pPr = paragraph._element.find(f"{{{W}}}pPr")
    if pPr is not None:
        outline = pPr.find(f"{{{W}}}outlineLvl")
        if outline is not None:
            val = outline.get(f"{{{W}}}val")
            if val is not None and val.isdigit():
                return int(val) + 1  # outlineLvl 0 = Heading 1

    return None


def _build_heading_tree(
    headings: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build a nested tree from a flat heading list using a stack.

    Each node: {level, text, paragraph_index, children: [...]}.
    """
    root: list[dict[str, Any]] = []
    stack: list[dict[str, Any]] = []  # tracks nesting context

    for h in headings:
        node = {
            "level": h["level"],
            "text": h["text"],
            "paragraph_index": h["paragraph_index"],
            "children": [],
        }

        # Pop stack until we find a parent with a lower level
        while stack and stack[-1]["level"] >= node["level"]:
            stack.pop()

        if stack:
            stack[-1]["children"].append(node)
        else:
            root.append(node)

        stack.append(node)

    return root


def _section_paragraphs(
    paragraphs, heading_idx: int, heading_level: int
) -> list:
    """Return paragraphs belonging to a section (from heading to next
    same-or-higher-level heading, exclusive)."""
    result = []
    for i in range(heading_idx + 1, len(paragraphs)):
        lvl = _detect_heading_level(paragraphs[i])
        if lvl is not None and lvl <= heading_level:
            break
        result.append(paragraphs[i])
    return result


# ── public API ─────────────────────────────────────────────────────


def load_document_structure(filepath: str) -> dict[str, Any]:
    """Parse a .docx and return heading tree + metadata.

    Returns:
      {
        filepath, paragraph_count, heading_count,
        headings: [{level, text, paragraph_index}, ...],
        heading_tree: [{level, text, paragraph_index, children}, ...]
      }
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    paragraphs = doc.paragraphs

    headings = []
    for idx, p in enumerate(paragraphs):
        level = _detect_heading_level(p)
        if level is not None:
            headings.append(
                {"level": level, "text": p.text.strip(), "paragraph_index": idx}
            )

    return {
        "filepath": filepath,
        "paragraph_count": len(paragraphs),
        "heading_count": len(headings),
        "headings": headings,
        "heading_tree": _build_heading_tree(headings),
    }


def get_section_content(filepath: str, heading_text: str) -> dict[str, Any]:
    """Return full text content under a specific heading.

    Collects all paragraphs from the heading until the next heading
    of the same or higher level.

    Returns:
      {heading, level, content, paragraph_count, subsections}
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    paragraphs = doc.paragraphs

    # Find the target heading
    target_idx = None
    target_level = None
    for idx, p in enumerate(paragraphs):
        level = _detect_heading_level(p)
        if level is not None and p.text.strip() == heading_text:
            target_idx = idx
            target_level = level
            break

    if target_idx is None:
        return {"error": f"Heading not found: {heading_text}"}

    section_pars = _section_paragraphs(paragraphs, target_idx, target_level)
    content_lines = [p.text for p in section_pars]
    content = "\n".join(content_lines)

    # Identify direct sub-headings
    subsections = []
    for p in section_pars:
        lvl = _detect_heading_level(p)
        if lvl is not None:
            subsections.append(p.text.strip())

    return {
        "heading": heading_text,
        "level": target_level,
        "content": content,
        "paragraph_count": len(section_pars),
        "subsections": subsections,
    }


def get_all_sections_summary(filepath: str) -> dict[str, Any]:
    """Return a compact summary for each section: preview + content hash.

    Used for quick redundancy screening — identical content_hash means
    exact duplicate, similar previews suggest near-duplicates.

    Returns:
      {filepath, sections: [{heading, level, preview, content_hash,
                              paragraph_count}, ...]}
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    paragraphs = doc.paragraphs

    # Collect all headings with their indices and levels
    headings = []
    for idx, p in enumerate(paragraphs):
        level = _detect_heading_level(p)
        if level is not None:
            headings.append((idx, level, p.text.strip()))

    sections = []
    for i, (h_idx, h_level, h_text) in enumerate(headings):
        pars = _section_paragraphs(paragraphs, h_idx, h_level)
        content = "\n".join(p.text for p in pars)

        preview = content[:200]
        if len(content) > 200:
            preview += "…"

        content_hash = hashlib.md5(content.encode("utf-8")).hexdigest()

        sections.append(
            {
                "heading": h_text,
                "level": h_level,
                "preview": preview,
                "content_hash": content_hash,
                "paragraph_count": len(pars),
            }
        )

    return {"filepath": filepath, "sections": sections}
