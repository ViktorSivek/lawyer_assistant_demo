#!/usr/bin/env python3
"""Generate a sample Czech legal .docx with intentional issues for testing.

Embedded issues:
  REDUNDANCY:
    1. Article 6 ¶1 = near-exact copy of Article 2 ¶1 (scope restated)
    2. Article 9 ¶3 = verbatim copy of Article 7 ¶1 (penalty clause duplicated)
    3. Příloha 2 last ¶ ≈ Article 5.1 ¶1 (near-duplicate, slight wording change)
    4. Příloha 1 is NOT redundant — it elaborates Article 2 (should NOT be flagged)

  CROSS-REFERENCES:
    Valid:  čl. 4, čl. 7, článek 4, článek 6, článek 3, příloha č. 1/2, § 2586
    Invalid: článek 12 (doesn't exist), příloha č. 5 (doesn't exist)
    All are plain text (none use field codes → should be flagged as format issue)

  WHITESPACE:
    1. Double spaces: Article 3 ¶1, Article 10 ¶3, Příloha 3 last ¶
    2. Trailing whitespace: Article 4 ¶2
    3. Leading whitespace: Article 7 ¶3
    4. Consecutive blank paragraphs: Article 6 between ¶2 and ¶3

  ENUMERATIONS:
    1. Good list: Article 2.1 — consistent semicolons, period at end
    2. Bad list:  Article 3.1 — mixed delimiters (comma/semicolon/none)
    3. Bad list:  Article 8 — last items inconsistent (comma then nothing)
    4. Good list: Příloha 2 — consistent semicolons, period at end
"""

import os
import sys

from docx import Document
from docx.shared import Pt
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _qn(tag: str) -> str:
    """Qualified name helper for OOXML tags."""
    return f"{{{W}}}{tag}"


def add_bookmark(paragraph, name: str, bm_id: int):
    """Insert a bookmark spanning the whole paragraph."""
    p_elem = paragraph._element

    bm_start = etree.SubElement(p_elem, _qn("bookmarkStart"))
    bm_start.set(_qn("id"), str(bm_id))
    bm_start.set(_qn("name"), name)

    bm_end = etree.SubElement(p_elem, _qn("bookmarkEnd"))
    bm_end.set(_qn("id"), str(bm_id))


def create_test_document(output_path: str):
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    bm_id = 0  # bookmark counter

    # ── TITLE ──────────────────────────────────────────────
    doc.add_heading("Smlouva o dílo č. 2024/001", level=0)

    # ── ČLÁNEK 1 ───────────────────────────────────────────
    h = doc.add_heading("Článek 1 – Smluvní strany", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_1", bm_id)

    doc.add_paragraph(
        "Objednatel: Město Příbram, IČO: 00243132, se sídlem Tyršova 108, "
        "261 01 Příbram I, zastoupené starostou Ing. Janem Konvalinkou "
        '(dále jen „Objednatel").'
    )
    doc.add_paragraph(
        "Zhotovitel: ABC Stavby s.r.o., IČO: 12345678, se sídlem Pražská 15, "
        "261 01 Příbram, zastoupená jednatelem Petrem Dvořákem "
        '(dále jen „Zhotovitel").'
    )

    # ── ČLÁNEK 2 ───────────────────────────────────────────
    h = doc.add_heading("Článek 2 – Předmět smlouvy", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_2", bm_id)

    doc.add_paragraph(
        "Předmětem této smlouvy je provedení stavebních prací na objektu "
        "Základní školy Příbram, ul. Školní 5, a to v rozsahu dle projektové "
        "dokumentace, která tvoří přílohu č. 1 této smlouvy "
        '(dále jen „Dílo").'
    )

    doc.add_heading("2.1 Rozsah díla", level=2)
    doc.add_paragraph("Dílo zahrnuje zejména:")

    # ✓ GOOD enumeration — consistent semicolons, period at end
    doc.add_paragraph(
        "a) demoliční práce dle bodu 3.1 projektové dokumentace;",
        style="List Number",
    )
    doc.add_paragraph(
        "b) stavební úpravy nosných konstrukcí;", style="List Number"
    )
    doc.add_paragraph(
        "c) instalace nových rozvodů elektřiny a vody;", style="List Number"
    )
    doc.add_paragraph(
        "d) dokončovací a úklidové práce.", style="List Number"
    )

    doc.add_heading("2.2 Místo plnění", level=2)
    doc.add_paragraph(
        "Místem plnění je objekt Základní školy Příbram na adrese "
        "Školní 5, 261 01 Příbram."
    )

    # ── ČLÁNEK 3 ───────────────────────────────────────────
    h = doc.add_heading("Článek 3 – Cena díla", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_3", bm_id)

    # ✗ WHITESPACE: double spaces
    doc.add_paragraph(
        "Celková cena díla  činí 4 500 000 Kč bez DPH  (slovy: čtyři miliony "
        "pět set tisíc korun českých). Cena je stanovena jako cena nejvýše "
        "přípustná."
    )

    doc.add_heading("3.1 Platební podmínky", level=2)
    doc.add_paragraph("Cena bude hrazena následovně:")

    # ✗ BAD enumeration — mixed delimiters: comma, semicolon, comma, nothing
    doc.add_paragraph(
        "(i) záloha ve výši 30 % při podpisu smlouvy,"
    )
    doc.add_paragraph(
        "(ii) průběžné měsíční fakturace dle skutečně provedených prací;"
    )
    doc.add_paragraph("(iii) závěrečná faktura po předání díla,")
    doc.add_paragraph(
        "(iv) pozastávka ve výši 10 % bude uvolněna po uplynutí záruční doby"
    )

    doc.add_heading("3.2 Fakturace", level=2)
    doc.add_paragraph(
        "Zhotovitel je oprávněn fakturovat provedené práce měsíčně, vždy "
        "k poslednímu dni kalendářního měsíce. Splatnost faktur činí 30 dnů "
        "ode dne doručení Objednateli."
    )

    # ── ČLÁNEK 4 ───────────────────────────────────────────
    h = doc.add_heading("Článek 4 – Termín plnění", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_4", bm_id)

    doc.add_paragraph(
        "Zhotovitel se zavazuje provést dílo v následujících termínech:"
    )
    doc.add_paragraph(
        "a) zahájení prací: do 14 dnů od nabytí účinnosti smlouvy;"
    )
    doc.add_paragraph("b) dokončení díla: nejpozději do 31. 12. 2025.")

    # ✗ WHITESPACE: trailing spaces
    doc.add_paragraph(
        "V případě prodlení Zhotovitele s dokončením díla dle článku 4 "
        "je Objednatel oprávněn požadovat smluvní pokutu dle čl. 7 "
        "této smlouvy.   "
    )

    # ── ČLÁNEK 5 ───────────────────────────────────────────
    h = doc.add_heading("Článek 5 – Práva a povinnosti smluvních stran", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_5", bm_id)

    doc.add_heading("5.1 Povinnosti Zhotovitele", level=2)
    doc.add_paragraph(
        "Zhotovitel je povinen provést dílo řádně, v souladu s touto smlouvou, "
        "projektovou dokumentací a platnými technickými normami. Zhotovitel je "
        "dále povinen dodržovat bezpečnostní předpisy dle přílohy č. 2."
    )
    # plain-text references (should be field codes)
    doc.add_paragraph(
        "Zhotovitel je povinen postupovat v souladu s harmonogramem dle "
        "článku 4 a dodržet cenový limit stanovený v článku 3."
    )

    doc.add_heading("5.2 Povinnosti Objednatele", level=2)
    doc.add_paragraph(
        "Objednatel je povinen předat Zhotoviteli staveniště do 7 dnů od "
        "nabytí účinnosti smlouvy. Objednatel je dále povinen poskytnout "
        "součinnost potřebnou pro řádné provedení díla."
    )

    # ── ČLÁNEK 6 ───────────────────────────────────────────
    h = doc.add_heading("Článek 6 – Předání a převzetí díla", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_6", bm_id)

    # ✗ REDUNDANCY #1: near-exact copy of Article 2 ¶1
    doc.add_paragraph(
        "Předmětem předání je provedení stavebních prací na objektu "
        "Základní školy Příbram, ul. Školní 5, a to v rozsahu dle "
        "projektové dokumentace, která tvoří přílohu č. 1 této smlouvy."
    )

    doc.add_paragraph(
        "O předání a převzetí díla bude sepsán předávací protokol podepsaný "
        "oběma smluvními stranami. Dílo se považuje za předané okamžikem "
        "podpisu předávacího protokolu."
    )

    # ✗ WHITESPACE: consecutive blank paragraphs
    doc.add_paragraph("")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Objednatel je oprávněn odmítnout převzetí díla, pokud dílo vykazuje "
        "vady bránící jeho řádnému užívání."
    )

    # ── ČLÁNEK 7 ───────────────────────────────────────────
    h = doc.add_heading("Článek 7 – Smluvní pokuty", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_7", bm_id)

    doc.add_paragraph(
        "V případě prodlení Zhotovitele s termínem dokončení díla dle čl. 4 "
        "je Objednatel oprávněn požadovat smluvní pokutu ve výši 0,05 % "
        "z celkové ceny díla za každý započatý den prodlení."
    )

    # ✗ INVALID REFERENCE: článek 12 doesn't exist
    doc.add_paragraph(
        "V případě porušení povinností dle článku 12 je Zhotovitel povinen "
        "uhradit smluvní pokutu ve výši 50 000 Kč za každý jednotlivý "
        "případ porušení."
    )

    # ✗ WHITESPACE: leading whitespace
    doc.add_paragraph(
        "  Uplatněním smluvní pokuty není dotčeno právo na náhradu škody."
    )

    # ── ČLÁNEK 8 ───────────────────────────────────────────
    h = doc.add_heading("Článek 8 – Záruční podmínky", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_8", bm_id)

    doc.add_paragraph(
        "Zhotovitel poskytuje na dílo záruku v délce 60 měsíců ode dne "
        "předání a převzetí díla dle článku 6 této smlouvy."
    )

    doc.add_paragraph("Záruční podmínky se nevztahují na:")

    # ✗ BAD enumeration — last items: semicolon, semicolon, comma, nothing
    doc.add_paragraph(
        "(a) vady způsobené nesprávným užíváním díla Objednatelem;"
    )
    doc.add_paragraph(
        "(b) vady vzniklé v důsledku zásahu vyšší moci;"
    )
    doc.add_paragraph("(c) běžné opotřebení díla,")
    doc.add_paragraph(
        "(d) vady způsobené zásahem třetích osob bez souhlasu Zhotovitele"
    )

    # ── ČLÁNEK 9 ───────────────────────────────────────────
    h = doc.add_heading("Článek 9 – Odstoupení od smlouvy", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_9", bm_id)

    doc.add_paragraph(
        "Objednatel je oprávněn odstoupit od smlouvy v případě, že Zhotovitel "
        "je v prodlení s dokončením díla o více než 30 dnů oproti termínu "
        "dle čl. 4."
    )
    doc.add_paragraph(
        "Zhotovitel je oprávněn odstoupit od smlouvy v případě, že Objednatel "
        "je v prodlení s úhradou faktur o více než 60 dnů."
    )

    # ✗ REDUNDANCY #2: verbatim copy of Article 7 ¶1
    doc.add_paragraph(
        "V případě prodlení Zhotovitele s termínem dokončení díla dle čl. 4 "
        "je Objednatel oprávněn požadovat smluvní pokutu ve výši 0,05 % "
        "z celkové ceny díla za každý započatý den prodlení."
    )

    # ── ČLÁNEK 10 ──────────────────────────────────────────
    h = doc.add_heading("Článek 10 – Závěrečná ustanovení", level=1)
    bm_id += 1
    add_bookmark(h, "clanek_10", bm_id)

    doc.add_paragraph(
        "Tato smlouva se řídí zákonem č. 89/2012 Sb., občanský zákoník, "
        "v platném znění. Smlouva nabývá účinnosti dnem podpisu oběma "
        "smluvními stranami."
    )

    # reference to §
    doc.add_paragraph(
        "Pro účely této smlouvy se použijí ustanovení § 2586 a násl. "
        "občanského zákoníku o smlouvě o dílo."
    )

    # ✗ WHITESPACE: double spaces
    doc.add_paragraph(
        "Smlouva je vyhotovena ve  dvou stejnopisech, z nichž  každá "
        "smluvní strana obdrží po jednom výtisku."
    )

    doc.add_paragraph(
        "Nedílnou součástí této smlouvy jsou následující přílohy:"
    )
    doc.add_paragraph("Příloha č. 1 – Projektová dokumentace")
    doc.add_paragraph("Příloha č. 2 – Bezpečnostní předpisy")
    doc.add_paragraph("Příloha č. 3 – Harmonogram prací")

    # ✗ INVALID REFERENCE: příloha č. 5 doesn't exist
    doc.add_paragraph(
        "Podrobnosti o pojištění jsou uvedeny v příloze č. 5 této smlouvy."
    )

    # ── PŘÍLOHA 1 ──────────────────────────────────────────
    doc.add_page_break()
    h = doc.add_heading("Příloha č. 1 – Projektová dokumentace", level=1)
    bm_id += 1
    add_bookmark(h, "priloha_1", bm_id)

    # Elaboration of Article 2 — NOT redundant
    doc.add_paragraph(
        "Projektová dokumentace pro provedení stavebních prací na objektu "
        "Základní školy Příbram specifikuje následující rozsah prací:"
    )

    doc.add_heading("1.1 Demoliční práce", level=2)
    doc.add_paragraph(
        "Demolice stávajících příček v 1. NP a 2. NP budovy, odstranění "
        "podlahových krytin a demontáž stávajících rozvodů elektřiny. "
        "Celkový rozsah demoličních prací: cca 450 m² podlahové plochy."
    )

    doc.add_heading("1.2 Stavební úpravy", level=2)
    doc.add_paragraph(
        "Vyzdění nových příček dle výkresové dokumentace, provedení nových "
        "podlah, oprava fasády v rozsahu dle výkresu č. D.1.4."
    )

    doc.add_heading("1.3 Instalace rozvodů", level=2)
    doc.add_paragraph(
        "Nové rozvody elektřiny (silnoproud i slaboproud), vody a kanalizace "
        "dle jednotlivých profesních částí projektové dokumentace."
    )

    # ── PŘÍLOHA 2 ──────────────────────────────────────────
    doc.add_page_break()
    h = doc.add_heading("Příloha č. 2 – Bezpečnostní předpisy", level=1)
    bm_id += 1
    add_bookmark(h, "priloha_2", bm_id)

    doc.add_paragraph(
        "Zhotovitel je povinen dodržovat následující bezpečnostní předpisy "
        "při provádění díla:"
    )

    # ✓ GOOD enumeration — consistent semicolons, period at end
    doc.add_paragraph(
        "(a) Zákon č. 309/2006 Sb., o zajištění dalších podmínek bezpečnosti "
        "a ochrany zdraví při práci;"
    )
    doc.add_paragraph(
        "(b) Nařízení vlády č. 591/2006 Sb., o bližších minimálních "
        "požadavcích na bezpečnost a ochranu zdraví při práci na "
        "staveništích;"
    )
    doc.add_paragraph(
        "(c) Nařízení vlády č. 362/2005 Sb., o bližších požadavcích na "
        "bezpečnost a ochranu zdraví při práci na pracovištích s nebezpečím "
        "pádu z výšky."
    )

    # ✗ REDUNDANCY #3: near-duplicate of Article 5.1 ¶1 (slight wording change)
    doc.add_paragraph(
        "Zhotovitel je povinen provést dílo řádně, v souladu s touto smlouvou, "
        "projektovou dokumentací a platnými technickými normami. Zhotovitel "
        "je povinen dodržovat bezpečnostní předpisy dle této přílohy."
    )

    # ── PŘÍLOHA 3 ──────────────────────────────────────────
    doc.add_page_break()
    h = doc.add_heading("Příloha č. 3 – Harmonogram prací", level=1)
    bm_id += 1
    add_bookmark(h, "priloha_3", bm_id)

    doc.add_paragraph(
        "Harmonogram prací je stanoven v souladu s článkem 4 této smlouvy."
    )

    doc.add_paragraph(
        "Etapa 1 (měsíc 1–2): Demoliční práce a příprava staveniště"
    )
    doc.add_paragraph(
        "Etapa 2 (měsíc 2–4): Stavební úpravy nosných konstrukcí"
    )
    doc.add_paragraph("Etapa 3 (měsíc 4–6): Instalace rozvodů")
    doc.add_paragraph("Etapa 4 (měsíc 6–7): Dokončovací práce a úklid")

    # ✗ WHITESPACE: double space
    doc.add_paragraph(
        "Zhotovitel je povinen  informovat Objednatele o průběhu prací "
        "minimálně jednou týdně."
    )

    # ── save ───────────────────────────────────────────────
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "documents",
        "test_smlouva.docx",
    )
    create_test_document(path)
